# -*- coding: utf-8 -*-
"""공통 도우미 : LGLS WCS 문서 산출물 생성용 (2026-09-04 판 : 숨김 메뉴·시뮬레이터 표기 제외)"""
import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'D:\project\LGLS\Renewal\docs\산출물_20260904'
SHOTS = r'C:\Users\USER\AppData\Local\Temp\claude\D--project-LGLS-Renewal\4609e0ee-85bd-49e7-8c4e-b9231f6fb7f0\scratchpad\shots_0904'
SCRATCH = r'C:\Users\USER\AppData\Local\Temp\claude\D--project-LGLS-Renewal\4609e0ee-85bd-49e7-8c4e-b9231f6fb7f0\scratchpad'
os.makedirs(OUT, exist_ok=True)

NAVY = RGBColor(0x1E, 0x27, 0x61)
GRAY = RGBColor(0x66, 0x66, 0x66)
DATE = '2026-09-04'
AUTHOR = 'LGLS WCS Renewal'

def new_doc(title, subtitle=None, landscape=False):
    d = Document()
    st = d.styles['Normal']; st.font.name = '맑은 고딕'; st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    for lv, sz in ((1, 16), (2, 13), (3, 11)):
        h = d.styles['Heading %d' % lv]; h.font.name = '맑은 고딕'; h.font.size = Pt(sz); h.font.bold = True
        h.font.color.rgb = NAVY; h.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    sec = d.sections[0]
    if landscape:
        from docx.enum.section import WD_ORIENT
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = sec.right_margin = Cm(2.0); sec.top_margin = sec.bottom_margin = Cm(2.0)
    for _ in range(6): d.add_paragraph()
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LG 화학 1동 자동창고'); r.font.size = Pt(16); r.font.color.rgb = GRAY
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.font.size = Pt(28); r.bold = True; r.font.color.rgb = NAVY
    if subtitle:
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle); r.font.size = Pt(13); r.font.color.rgb = GRAY
    for _ in range(10): d.add_paragraph()
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('WCS Renewal (구 ECS 대체)   |   작성일 %s   |   Ver 1.1' % DATE); r.font.size = Pt(10); r.font.color.rgb = GRAY
    d.add_page_break()
    d.add_heading('문서 이력', 1)
    table(d, ['버전', '일자', '작성', '내용'], [['1.0', '2026-09-03', AUTHOR, '최초 작성'], ['1.1', DATE, AUTHOR, '기타 수정']], widths=[1.5, 2.5, 3.0, 9.0])
    d.add_page_break()
    return d

def shade(cell, hex_):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hex_)
    tcPr.append(sh)

def table(d, headers, rows, widths=None, font=9, header_fill='1E2761'):
    t = d.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(str(h)); r.bold = True; r.font.size = Pt(font); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; shade(c, header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run('' if v is None else str(v)); r.font.size = Pt(font)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    d.add_paragraph()
    return t

def para(d, text, bold=False, size=10, color=None, italic=False, indent=0):
    p = d.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def bullets(d, items, indent=0.5):
    for it in items:
        p = d.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent = Cm(indent + 0.6)
        r = p.add_run(it); r.font.size = Pt(10)

def numbered(d, items):
    for it in items:
        p = d.add_paragraph(style='List Number'); r = p.add_run(it); r.font.size = Pt(10)

def image(d, path, width_cm=16.0, caption=None):
    if not os.path.exists(path):
        para(d, '[화면 캡처 없음: %s]' % os.path.basename(path), italic=True, color=GRAY); return
    d.add_picture(path, width=Cm(width_cm))
    d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption); r.font.size = Pt(9); r.font.color.rgb = GRAY; r.italic = True

def shot(name):
    return os.path.join(SHOTS, name)

def read_cp949(path):
    return open(path, 'rb').read().decode('cp949', errors='replace')

def menu_states():
    ini = read_cp949(r'D:\project\LGLS\Renewal\WCS\CPlusPlus\LGLS_CLIENT\Bin\Debug\Ecs.ini')
    m = {}
    sec = re.search(r'\[MENU\](.*?)(?:\r?\n\[|\Z)', ini, re.S)
    if sec:
        for line in sec.group(1).splitlines():
            mm = re.match(r'\s*([A-Z_]+)\s*=\s*(\d)', line)
            if mm: m[mm.group(1)] = mm.group(2) == '1'
    return m
