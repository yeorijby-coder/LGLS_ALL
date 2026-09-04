# -*- coding: utf-8 -*-
"""이중입고(54) / 공출고(58) 반복 시험 보고서 (docx)"""
import os, re, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'D:\project\LGLS\Renewal\docs\산출물_20260904'
RESULT = r'C:\Users\USER\AppData\Local\Temp\claude\D--project-LGLS-Renewal\4609e0ee-85bd-49e7-8c4e-b9231f6fb7f0\scratchpad\err_test_result.txt'
NAVY = RGBColor(0x1E, 0x27, 0x61)
GRAY = RGBColor(0x66, 0x66, 0x66)
DATE = '2026-09-05'


def shade(cell, hex_):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hex_)
    tcPr.append(sh)


def table(d, headers, rows, widths=None, font=9):
    t = d.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(str(h)); r.bold = True
        r.font.size = Pt(font); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; shade(c, '1E2761')
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run('' if v is None else str(v))
            r.font.size = Pt(font)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    d.add_paragraph()
    return t


def para(d, text, bold=False, size=10, color=None, indent=0):
    p = d.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p


def bullets(d, items, indent=0.4):
    for it in items:
        p = d.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(indent + 0.6)
        r = p.add_run(it); r.font.size = Pt(10)


# ── 결과 파일 파싱 ───────────────────────────────────────────
txt = open(RESULT, encoding='utf-8-sig').read()
rounds = {'DOUBLE': [], 'EMPTY': []}
cur = None
for ln in txt.splitlines():
    if '이중입고(54) 시험' in ln: cur = 'DOUBLE'; continue
    if '공출고(58) 시험' in ln: cur = 'EMPTY'; continue
    m = re.search(r'#(\d+)\s+(OK|NG)\s*:\s*(.+)$', ln)
    if m and cur:
        rounds[cur].append((int(m.group(1)), m.group(2), m.group(3).strip(),
                            ln.split()[0]))
okD = sum(1 for r in rounds['DOUBLE'] if r[1] == 'OK')
okE = sum(1 for r in rounds['EMPTY'] if r[1] == 'OK')
ngD = len(rounds['DOUBLE']) - okD
ngE = len(rounds['EMPTY']) - okE

# ── 문서 ─────────────────────────────────────────────────────
d = Document()
st = d.styles['Normal']; st.font.name = '맑은 고딕'; st.font.size = Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
for lv, sz in ((1, 15), (2, 12)):
    h = d.styles['Heading %d' % lv]; h.font.name = '맑은 고딕'
    h.font.size = Pt(sz); h.font.bold = True; h.font.color.rgb = NAVY
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
sec = d.sections[0]
sec.left_margin = sec.right_margin = Cm(2.0)
sec.top_margin = sec.bottom_margin = Cm(1.8)

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('LG 화학 1동 자동창고'); r.font.size = Pt(12); r.font.color.rgb = GRAY
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('이중입고 · 공출고 반복 시험 보고서'); r.font.size = Pt(20)
r.bold = True; r.font.color.rgb = NAVY
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('시험일 %s   ·   WCS Renewal' % DATE); r.font.size = Pt(10)
r.font.color.rgb = GRAY
d.add_paragraph()

d.add_heading('1. 시험 목적과 방법', 1)
para(d, '크레인이 이중입고(에러 54) 또는 공출고(에러 58)를 낸 뒤, 상위(WMS/IMS)의 재지정을 받아 '
        '작업이 정상으로 돌아오는지 반복 확인한다.')
table(d, ['구분', '내용'], [
    ['에러 주입', 'EQP_SIM 의 [이중입고 에러] / [공출고 에러] 체크박스'],
    ['에러 보고', 'WCS 가 E 전문으로 상위에 보고 (DeviceClass=1, ErrorKind 1=이중입고 / 3=공출고)'],
    ['재지정', 'HOST_SIM 의 [재지정 처리] 버튼 → R 전문 (R_Kind 1=입고 / 2=출고)'],
    ['성공 판정', '재지정 후 크레인 에러코드가 해제되고 작업이 다시 진행되면 성공'],
    ['반복 횟수', '각 12 회'],
], widths=[3.2, 13.5])

d.add_heading('2. 시험 결과', 1)
table(d, ['시험', '성공', '실패', '판정'], [
    ['이중입고 (에러 54)', okD, ngD, '합격' if ngD == 0 and okD >= 10 else '조건부'],
    ['공출고 (에러 58)', okE, ngE, '합격' if ngE == 0 and okE >= 10 else '조건부'],
], widths=[6.0, 2.5, 2.5, 5.7], font=10)

for key, title in (('DOUBLE', '2.1 이중입고 (에러 54)'), ('EMPTY', '2.2 공출고 (에러 58)')):
    d.add_heading(title, 2)
    rows = [[n, '성공' if st_ == 'OK' else '실패', tm, msg] for n, st_, msg, tm in rounds[key]]
    table(d, ['회차', '결과', '시각', '내용'], rows, widths=[1.6, 2.0, 2.6, 10.6])

d.add_heading('3. 시험 중 발견해 수정한 결함', 1)
para(d, '첫 시행에서는 두 시험 모두 전부 실패했다. 원인을 추적해 다음 여섯 가지를 고쳤다.')
table(d, ['No', '증상', '원인', '조치'], [
    ['1', '재지정을 보내도 작업이 에러 상태에 머문다',
     '스케줄러가 매 주기 에러 상태(08/07)로 되돌려, 재지정 상태(06/05)가 유지되지 않았다',
     'MarkErrorJobStatus 가 재지정 상태를 덮어쓰지 않도록 제외 조건 추가'],
    ['2', '재지정 상태가 되어도 크레인이 움직이지 않는다',
     '지시 기록이 "에러 없을 때만" 조건에 막혔다. 크레인 에러는 새 지시를 받아야 풀리므로 서로 물린다',
     '재지정 지시는 이중입고/공출고 에러가 남아 있어도 기록하도록 예외 허용'],
    ['3', '출고 재지정이 항상 실패',
     '존재하지 않는 테이블(LUGG_MST)과 컬럼(LUGGNO)을 갱신하고 있었다',
     'JOB_MST / LUGG_NO 로 정정'],
    ['4', '상위가 재지정을 보내면 "그런 작업 없음"으로 반려',
     '에러보고(E)의 작업번호를 LUGG_NO_FK1_RD 에서 읽는데, 실경로에서는 채워지지 않아 낡은 값이 올라갔다',
     '실린 화물(PALLET_ON_VEHICLE_RD → ITN_LUGG_FK1 → LUGG_NO_FK1_OD) 기준으로 정정'],
    ['5', '재지정이 "동일 호기가 아님"으로 반려',
     '에러보고의 셀이 크레인 현재 위치라 작업의 목적 셀과 달랐다',
     '이중입고는 작업의 도착 셀, 공출고는 출발 셀을 보고하도록 정정'],
    ['6', '공출고 재지정이 "기존과 같은 위치"로 반려',
     '새 셀을 전문의 도착 필드에 실었으나, 출고 재지정은 출발 필드가 바뀌어야 한다',
     'HOST_SIM 이 새 출발 셀을 START 필드에 싣도록 정정. 아울러 종전에는 R 전문 자체를 보내지 않았다'],
], widths=[1.0, 4.4, 6.0, 5.4], font=8.5)

d.add_heading('4. 정상 복구 흐름 (확인된 순서)', 1)
para(d, '이중입고', bold=True)
bullets(d, [
    '크레인이 목적 셀에 넣지 못하고 에러 54 발생 (화물은 실은 채)',
    'WCS 가 E 전문 보고 (ErrorKind=1, 작업번호와 도착 셀 포함)',
    '상위가 같은 호기의 다른 셀로 R 전문(R_Kind=1) 재지정',
    'WCS 가 작업 상태를 06(이중입고 재지정)으로 바꾸고, 스케줄러가 새 셀로 크레인 재지시',
    '크레인이 새 지시를 받으면서 에러 해제 → 정상 입고 완료',
])
para(d, '공출고', bold=True)
bullets(d, [
    '크레인이 출발 셀에서 집지 못하고 에러 58 발생 (빈 포크)',
    'WCS 가 E 전문 보고 (ErrorKind=3, 작업번호와 출발 셀 포함)',
    '상위가 같은 호기의 다른 셀을 출발지로 R 전문(R_Kind=2) 재지정',
    'WCS 가 작업 상태를 05(공출고 재지정)으로 바꾸고, 스케줄러가 새 셀로 크레인 재지시',
    '크레인이 새 지시를 받으면서 에러 해제 → 정상 출고 완료',
])

d.add_heading('5. 현장 확인이 필요한 사항', 1)
bullets(d, [
    '실 크레인이 에러 상태에서 새 지시를 받아 스스로 에러를 푸는지, 아니면 운전자의 에러 해제가 먼저 필요한지. '
    'EQP_SIM 은 새 지시 수신 시 해제하도록 맞춰 두었다.',
    '이중입고 재지정에서 실 크레인이 이미 든 화물을 다시 집지 않고 새 셀로 바로 옮기는지.',
    '재지정 전문의 셀 규약(동일 호기 안에서만 허용)이 현장 WMS 와 일치하는지.',
])
out = os.path.join(OUT, '08_이중입고_공출고_시험_보고서.docx')
d.save(out)
print('saved', out, '| DOUBLE %d/%d  EMPTY %d/%d' % (okD, okD + ngD, okE, okE + ngE))
